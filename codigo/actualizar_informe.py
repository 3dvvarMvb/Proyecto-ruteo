"""Actualiza Informe-Ruteo-Resiliente-Drones.docx:
- Reemplaza SENAPRED (Amenaza 2) por CONAF (ArcGIS REST publico, verificado).
- Ajusta tabla de taxonomia de amenazas, capas poligonales y APIs.
- Agrega limitaciones verificadas (Overpass caido, SENAPRED sin API publica).
"""
import copy
import docx

SRC = "/home/kali/ruteo/Informe-Ruteo-Resiliente-Drones.docx"
DST = "/home/kali/ruteo/Informe-Ruteo-Resiliente-Drones.docx"

d = docx.Document(SRC)


def set_para_text(paragraph, text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def set_cell_text(cell, text):
    set_para_text(cell.paragraphs[0], text)


def clone_row_with_texts(table, source_row_idx, new_texts):
    """Clona la fila source_row_idx (para heredar estilo/bordes) y la agrega
    al final de la tabla con nuevos textos."""
    src_tr = table.rows[source_row_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    table._tbl.append(new_tr)
    new_row = table.rows[-1]
    for cell, text in zip(new_row.cells, new_texts):
        set_cell_text(cell, text)
    return new_row


def clone_paragraph_after(paragraph, text):
    """Clona un parrafo (para heredar bullet/estilo) y lo inserta despues,
    con el texto nuevo."""
    new_p_elm = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p_elm)
    from docx.text.paragraph import Paragraph
    new_paragraph = Paragraph(new_p_elm, paragraph._parent)
    set_para_text(new_paragraph, text)
    return new_paragraph


# ---------------------------------------------------------------------
# 1) Encabezado Amenaza 2: SENAPRED -> CONAF
# ---------------------------------------------------------------------
set_para_text(
    d.paragraphs[147],
    "Amenaza 2 — CONAF: Pronóstico de Riesgo de Incendios (Integrante B)",
)

# ---------------------------------------------------------------------
# 2) Tabla 8 (Campo/Detalle de la Amenaza 2): contenido real verificado
# ---------------------------------------------------------------------
t8 = d.tables[8]
nuevo_contenido_t8 = [
    ("Qué entrega",
     "Polígonos de probabilidad de ignición a 5 días (pronóstico diario) y "
     "polígonos de Áreas Silvestres Protegidas, servidos por el dashboard "
     "ArcGIS Online público \"Pronóstico de Riesgo\" de CONAF/GEPRIF"),
    ("Aporte al modelo",
     "Cada polígono de alta probabilidad de ignición penaliza el costo de "
     "las aristas que intersecta (ST_Intersects); los polígonos de áreas "
     "protegidas alimentan además la capa de exclusión del §7.2"),
    ("Acceso",
     "REST/GeoJSON, sin API key ni registro (ArcGIS Feature Service de "
     "acceso público, verificado: access=\"public\")"),
    ("Endpoint — probabilidad de ignición",
     "https://services5.arcgis.com/A1ELWse9bRAi2JiV/arcgis/rest/services/PI/FeatureServer/4/query"),
    ("Endpoint — áreas protegidas",
     "https://services5.arcgis.com/A1ELWse9bRAi2JiV/arcgis/rest/services/ASP/FeatureServer/0/query"),
    ("Dashboard de origen",
     "https://www.conaf.cl/incendios/situacion-actual-y-pronostico-de-incendios/ "
     "(enlaza el dashboard ArcGIS, ítem público 06a31e138f5c40efbd577c1993154ce5)"),
    ("Evidencia",
     "Consulta GeoJSON real para el bbox de la Región de Valparaíso: 4 polígonos "
     "de riesgo de ignición y 13 polígonos de áreas protegidas devueltos con geometría"),
]
for row, (campo, detalle) in zip(t8.rows[1:], nuevo_contenido_t8):
    set_cell_text(row.cells[0], campo)
    set_cell_text(row.cells[1], detalle)

# ---------------------------------------------------------------------
# 3) Parrafo "Complemento util" -> ahora es SENAPRED el complemento
# ---------------------------------------------------------------------
set_para_text(
    d.paragraphs[149],
    "Nota sobre SENAPRED: su portal (https://senapred.cl) categoriza alertas "
    "por comuna (Roja/Amarilla/Preventiva), pero no expone una API pública — "
    "su backend real es un GraphQL (AWS AppSync) protegido con autenticación "
    "Cognito, verificado durante las pruebas de factibilidad. Se cita solo "
    "como referencia normativa/cualitativa, no como fuente programática.",
)

# ---------------------------------------------------------------------
# 4) Tabla 6 (taxonomia de amenazas): fila 2 -> riesgo de ignicion CONAF
# ---------------------------------------------------------------------
t6 = d.tables[6]
row2 = t6.rows[2]
set_cell_text(row2.cells[0], "Zona de alto riesgo de ignición (CONAF)")
set_cell_text(row2.cells[1], "Penalización regional gradual")
set_cell_text(row2.cells[2], "Incrementa el costo de todas las aristas dentro del polígono de riesgo (pronóstico a 1-5 días)")
set_cell_text(row2.cells[3], "Decenas de km")

# ---------------------------------------------------------------------
# 5) Tabla 11 (capas poligonales): actualizar "Limites comunales" y
#    agregar fila para el poligono de riesgo de ignicion CONAF
# ---------------------------------------------------------------------
t11 = d.tables[11]
for row in t11.rows[1:]:
    if row.cells[0].text.strip() == "Límites comunales":
        set_cell_text(row.cells[2], "Límites administrativos generales")
        break
clone_row_with_texts(
    t11, 1,
    ["Probabilidad de ignición (CONAF, 5 días)",
     "CONAF / ArcGIS Online (Feature Service público)",
     "Penalización dinámica de aristas por riesgo de incendio"],
)

# ---------------------------------------------------------------------
# 6) Tabla 10 (APIs utilizadas): agregar CONAF, 5 -> 6 fuentes
# ---------------------------------------------------------------------
t10 = d.tables[10]
clone_row_with_texts(
    t10, 1,
    ["CONAF (ArcGIS REST — Pronóstico de Riesgo)", "Ninguna", "Amenazas"],
)
for p in d.paragraphs:
    if p.text.startswith("Se utilizan cinco fuentes"):
        set_para_text(p, "Se utilizan seis fuentes con acceso programático, superando ampliamente el mínimo:")
        break

# ---------------------------------------------------------------------
# 7) Limitaciones (§10): agregar dos hallazgos verificados
# ---------------------------------------------------------------------
for i, p in enumerate(d.paragraphs):
    if p.text.startswith("Cobertura de OSM variable"):
        p1 = clone_paragraph_after(
            p,
            "Disponibilidad de Overpass API. Durante las pruebas de factibilidad "
            "se verificó una caída real del servicio público de Overpass (servidor "
            "principal y un espejo independiente, ambos inaccesibles simultáneamente). "
            "Como plan de contingencia para la presentación se usará el extracto "
            "offline de OpenStreetMap para Chile publicado por Geofabrik "
            "(download.geofabrik.de/south-america/chile-latest.osm.pbf, verificado accesible).",
        )
        clone_paragraph_after(
            p1,
            "SENAPRED no expone API pública. Se verificó que senapred.cl es una "
            "aplicación de una sola página cuyo backend (GraphQL vía AWS AppSync) "
            "exige autenticación de usuario (Cognito); no hay forma de consumirlo "
            "de forma programática sin credenciales. Por eso se reemplazó como "
            "fuente principal de la Amenaza 2 por el servicio REST público de CONAF.",
        )
        break

d.save(DST)
print("Informe actualizado y guardado en", DST)
