"""Correccion: el conteo de '56 focos en Chile' estaba mal. El bbox
rectangular de la consulta a FIRMS incluye territorio argentino/boliviano
(Chile es angosto), asi que se filtro por el poligono real del pais:
solo 7 de los 56 focos detectados estaban efectivamente en Chile."""
import docx

PATH = "/home/kali/ruteo/Informe-Ruteo-Resiliente-Drones.docx"
d = docx.Document(PATH)


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


t7 = d.tables[7]
for row in t7.rows[1:]:
    if row.cells[0].text.strip() == "Evidencia":
        set_cell_text(
            row.cells[1],
            "MAP_KEY real obtenida y probada (2026-08-26): 0 focos activos en la Región de "
            "Valparaíso en los últimos 5 días (esperable en invierno). Para confirmar que la "
            "consulta sí detecta incendios cuando existen, se probó con un bbox rectangular "
            "de todo Chile: dio 56 focos, pero al filtrar por el polígono real del país "
            "(Chile es angosto — un bbox rectangular incluye Argentina y Bolivia al este de "
            "la cordillera) solo 7 estaban efectivamente en territorio chileno. Esto confirma "
            "además un requisito de diseño para la Tarea 2: las consultas geoespaciales deben "
            "usar el polígono real de la zona de interés (ST_Intersects/ST_Within), no un "
            "bounding box, para no contaminar el modelo con eventos fuera del país.",
        )
        break

d.save(PATH)
print("Informe corregido (v3) y guardado en", PATH)
